from flask import Flask, render_template, request, session, jsonify, redirect, url_for, flash, send_file
import os
from transformers import pipeline
import logging
import json
from werkzeug.security import generate_password_hash, check_password_hash
import datetime
import uuid
from docx import Document
from docx.shared import Pt
import tempfile
import warnings  # Replace exceptions import with warnings
import re

# Try to import validators, or create a fallback implementation
try:
    import validators
except ImportError:
    # Create a simple email validation fallback
    class validators:
        @staticmethod
        def email(email):
            """Simple email validation as fallback."""
            import re
            pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
            return re.match(pattern, email) is not None

# Set environment variable to allow insecure transport for OAuth
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

app = Flask(__name__)
app.secret_key = os.urandom(24)  # Used for session encryption

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Initialize the LLM model for text generation (use a more capable model for legal questions)
llm = pipeline('text-generation', model='distilgpt2')

# Initialize a more capable model for legal questions
try:
    legal_llm = pipeline('text-generation', model='gpt2-medium')  # Better model than distilgpt2
except Exception as e:  # Fixed the syntax error here
    legal_llm = llm  # Fall back to distilgpt2 if gpt2-medium fails to load
    logging.error(f"Failed to load gpt2-medium: {e}")

# User database file path
USERS_DB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'users.json')

# Load or create users database
def load_users():
    if os.path.exists(USERS_DB_FILE):
        with open(USERS_DB_FILE, 'r') as f:
            return json.load(f)
    else:
        return {"users": {}}

def save_users(users):
    with open(USERS_DB_FILE, 'w') as f:
        json.dump(users, f)

@app.route("/")
def index():
    if 'user_email' in session:
        return redirect(url_for('dashboard'))
    return render_template("login.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        confirm_password = request.form.get("confirm_password", "")
        name = request.form.get("name", "").strip()
        
        # Server-side validation
        errors = []
        
        # Validate name (minimum 2 characters, only letters and spaces)
        if not re.match(r'^[A-Za-z\s]{2,}$', name):
            errors.append("Name must be at least 2 characters and contain only letters and spaces.")
        
        # Validate email - fallback implementation will be used if validators module is not installed
        if not validators.email(email):
            errors.append("Please provide a valid email address.")
        
        # Check if email already exists
        users = load_users()
        if email in users["users"]:
            errors.append("Email already registered! Please use a different email.")
        
        # Validate password complexity
        if len(password) < 8:
            errors.append("Password must be at least 8 characters.")
        if not re.search(r'[A-Z]', password):
            errors.append("Password must contain at least one uppercase letter.")
        if not re.search(r'[a-z]', password):
            errors.append("Password must contain at least one lowercase letter.")
        if not re.search(r'[0-9]', password):
            errors.append("Password must contain at least one number.")
        if not re.search(r'[^A-Za-z0-9]', password):
            errors.append("Password must contain at least one special character.")
            
        # Check if passwords match
        if password != confirm_password:
            errors.append("Passwords do not match.")
        
        # If there are validation errors, flash messages and return to register page
        if errors:
            for error in errors:
                flash(error, "danger")
            return render_template("register.html")
        
        # If validation passes, create the user
        users["users"][email] = {
            "password_hash": generate_password_hash(password),
            "name": name,
            "created_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        save_users(users)
        
        flash("Registration successful! Please login.", "success")
        return redirect(url_for("index"))
    
    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        
        users = load_users()
        
        if email not in users["users"]:
            flash("Invalid email or password!")
            return render_template("login.html")
        
        user = users["users"][email]
        
        if not check_password_hash(user["password_hash"], password):
            flash("Invalid email or password!")
            return render_template("login.html")
        
        # Add user to session
        session["user_email"] = email
        session["user_name"] = user["name"]
        
        return redirect(url_for("dashboard"))
    
    return render_template("login.html")

@app.route("/dashboard")
def dashboard():
    if 'user_email' not in session:
        return redirect(url_for('index'))
    
    if 'state' not in session:
        session['state'] = 'awaiting_mode_selection'
        session['data'] = {}
    
    user = {'displayName': session['user_name'], 'email': session['user_email']}
    return render_template("index.html", user=user)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('index'))

@app.route("/chat", methods=['POST'])
def chat():
    user_message = request.json.get('message', '').strip()
    response_message = ""
    
    state = session.get('state', 'awaiting_mode_selection')
    data = session.get('data', {})

    if state == 'awaiting_mode_selection':
        if "tenancy" in user_message.lower():
            data['document_type'] = "Tenancy Agreement"
            session['state'] = 'awaiting_landlord'
            response_message = ("Great! Let's draft your Indian Tenancy Agreement. "
                                "Please provide the full name of the landlord.")
        elif "contract" in user_message.lower():
            data['document_type'] = "Contract"
            session['state'] = 'awaiting_contract_type'
            response_message = ("Great! Let's draft your Indian Contract. "
                               "What type of contract do you need? (Examples: Service Agreement, "
                               "Employment Contract, Sale Agreement, Loan Agreement)")
        elif "legal question" in user_message.lower():
            data['document_type'] = "Legal Question"
            session['state'] = 'awaiting_question'
            response_message = ("Sure! Please ask your legal question about Indian law.")
        else:
            response_message = ("Please select a mode: 'tenancy agreement', 'contract', or 'legal question'.")
    
    elif state == 'awaiting_landlord':
        data['landlord'] = user_message
        session['state'] = 'awaiting_tenant'
        response_message = "Thank you. Now, please provide the full name of the tenant."
    
    elif state == 'awaiting_tenant':
        data['tenant'] = user_message
        session['state'] = 'awaiting_property_address'
        response_message = "Great! Now, please provide the complete property address."
    
    elif state == 'awaiting_property_address':
        data['property_address'] = user_message
        session['state'] = 'awaiting_rental_details'
        response_message = ("Got it. Please provide the monthly rent amount in INR and lease duration "
                            "(e.g., 'Rs. 15,000 per month for 11 months').")
    
    elif state == 'awaiting_rental_details':
        data['rental_details'] = user_message
        session['state'] = 'awaiting_security_deposit'
        response_message = "What is the security deposit amount in INR?"
    
    elif state == 'awaiting_security_deposit':
        data['security_deposit'] = user_message
        session['state'] = 'completed'
        draft = generate_tenancy_agreement(data)
        response_message = (
            "Here is your draft Indian Rental Agreement:\n\n" + draft + "\n\n"
            "Disclaimer: This draft is for informational purposes only and is not a substitute "
            "for professional legal advice. Please consult a legal professional for a final review."
        )
    
    elif state == 'awaiting_contract_type':
        data['contract_type'] = user_message
        session['state'] = 'awaiting_first_party'
        response_message = "Please provide the full name and address of the first party (usually the service provider, employer, seller, etc.)."
    
    elif state == 'awaiting_first_party':
        data['first_party'] = user_message
        session['state'] = 'awaiting_second_party'
        response_message = "Please provide the full name and address of the second party (usually the service receiver, employee, buyer, etc.)."
    
    elif state == 'awaiting_second_party':
        data['second_party'] = user_message
        session['state'] = 'awaiting_contract_value'
        response_message = "What is the total contract value or consideration amount (in INR)?"
    
    elif state == 'awaiting_contract_value':
        data['contract_value'] = user_message
        session['state'] = 'awaiting_contract_duration'
        response_message = "What is the duration of this contract? (Example: 12 months, 2 years, etc.)"
    
    elif state == 'awaiting_contract_duration':
        data['contract_duration'] = user_message
        session['state'] = 'awaiting_contract_terms'
        response_message = "Please provide any specific terms and conditions for this contract."
    
    elif state == 'awaiting_contract_terms':
        data['contract_terms'] = user_message
        session['state'] = 'completed'
        # Generate a unique ID for this contract
        contract_id = f"CNT-{uuid.uuid4().hex[:8].upper()}"
        data['contract_id'] = contract_id
        
        draft = generate_contract(data)
        
        # Save the contract to the user's session
        session['contract_data'] = data
        
        response_message = (
            f"Here is your draft Indian {data['contract_type']}:\n\n"
            f"{draft[:500]}...\n\n"  # Show just the beginning
            f"Your contract ID is: {contract_id}\n\n"
            f"You can <a href='/download_contract/{contract_id}' target='_blank'>download the full contract as a Word document</a>.\n\n"
            f"Disclaimer: This draft is for informational purposes only and is not a substitute "
            f"for professional legal advice. Please consult a legal professional for a final review."
        )
    
    elif state == 'awaiting_question':
        # Use legal_llm for legal questions
        prompt = f"Answer this Indian legal question accurately and professionally: {user_message}\nAnswer: "
        
        try:
            response = legal_llm(prompt, max_length=200, num_return_sequences=1, 
                              temperature=0.7, top_k=50, truncation=True)
            response_text = response[0]['generated_text']
            
            # Extract just the answer part
            if "Answer: " in response_text:
                response_message = response_text.split("Answer: ")[1].strip()
            else:
                response_message = response_text.replace(prompt, "").strip()
                
            # Add disclaimer
            response_message += "\n\nDisclaimer: This information is for general guidance only and should not be considered as legal advice."
            
        except Exception as e:
            response_message = f"I apologize, but I couldn't generate a response to your legal question. Please try asking in a different way."
            logging.error(f"Error generating legal response: {str(e)}")
        
        session['state'] = 'awaiting_mode_selection'
    
    elif state == 'completed':
        response_message = ("The process is complete. "
                            "If you need to start over, please select a mode: 'tenancy agreement', 'contract', or 'legal question'.")
        session['state'] = 'awaiting_mode_selection'
    
    else:
        response_message = ("I'm sorry, I didn't understand that. "
                            "Let's start over. What type of legal document do you need help with?")
        session['state'] = 'awaiting_mode_selection'
    
    session['data'] = data
    return jsonify({"response": response_message})

@app.route("/download_contract/<contract_id>")
def download_contract(contract_id):
    if 'user_email' not in session or 'contract_data' not in session:
        flash("Please generate a contract first.")
        return redirect(url_for('dashboard'))
    
    contract_data = session.get('contract_data', {})
    
    if contract_data.get('contract_id') != contract_id:
        flash("Contract not found.")
        return redirect(url_for('dashboard'))
    
    # Generate the Word document
    doc = Document()
    
    # Add heading
    heading = doc.add_heading(f"{contract_data.get('contract_type', 'CONTRACT AGREEMENT')}", 0)
    heading.alignment = 1  # Center alignment
    
    # Add subheading
    subheading = doc.add_paragraph("(As per Indian Contract Act, 1872)")
    subheading.alignment = 1  # Center alignment
    
    # Add contract number and date
    today = datetime.datetime.now().strftime("%d-%m-%Y")
    doc.add_paragraph(f"Contract No: {contract_id}")
    doc.add_paragraph(f"Date: {today}")
    
    # Add intro paragraph
    doc.add_paragraph(f"THIS CONTRACT AGREEMENT is made and executed at _______________ on this {today} between:")
    
    # Add parties
    doc.add_paragraph(f"{contract_data.get('first_party', 'First Party')}, hereinafter referred to as the \"FIRST PARTY\"")
    doc.add_paragraph("AND")
    doc.add_paragraph(f"{contract_data.get('second_party', 'Second Party')}, hereinafter referred to as the \"SECOND PARTY\"")
    
    # Add contract details
    doc.add_heading("WHEREAS:", level=1)
    doc.add_paragraph(f"The parties have agreed to enter into a {contract_data.get('contract_type', 'Contract')} as detailed below.")
    
    # Add terms and conditions
    doc.add_heading("NOW THIS AGREEMENT WITNESSES AS FOLLOWS:", level=1)
    
    # 1. Scope
    doc.add_heading("1. SCOPE OF WORK:", level=2)
    doc.add_paragraph(f"This {contract_data.get('contract_type', 'Contract')} covers the following services/goods/work:")
    
    # 2. Consideration
    doc.add_heading("2. CONSIDERATION:", level=2)
    doc.add_paragraph(f"The total consideration for this contract is {contract_data.get('contract_value', '________')}.")
    
    # 3. Term
    doc.add_heading("3. TERM:", level=2)
    doc.add_paragraph(f"This agreement shall be valid for a period of {contract_data.get('contract_duration', '________')}, commencing from {today}.")
    
    # 4. Specific terms
    doc.add_heading("4. SPECIFIC TERMS AND CONDITIONS:", level=2)
    doc.add_paragraph(f"{contract_data.get('contract_terms', 'Standard terms and conditions apply.')}")
    
    # 5-7. Standard clauses
    doc.add_heading("5. GOVERNING LAW:", level=2)
    doc.add_paragraph("This agreement shall be governed by and construed in accordance with the laws of India, and the courts in India shall have exclusive jurisdiction.")
    
    doc.add_heading("6. DISPUTE RESOLUTION:", level=2)
    doc.add_paragraph("Any dispute arising out of or in connection with this agreement shall first be resolved through amicable settlement. If not resolved within 30 days, the dispute shall be referred to arbitration in accordance with the Arbitration and Conciliation Act, 1996.")
    
    doc.add_heading("7. FORCE MAJEURE:", level=2)
    doc.add_paragraph("Neither party shall be liable for any failure or delay in performance due to circumstances beyond its reasonable control.")
    
    # Signature section
    doc.add_paragraph("\n\n")
    doc.add_paragraph("IN WITNESS WHEREOF, the parties hereto have executed this agreement on the day and year first above written.")
    doc.add_paragraph("\n\n")
    
    signature_table = doc.add_table(rows=1, cols=2)
    cell1 = signature_table.cell(0, 0)
    cell1.text = "FIRST PARTY:\n\n\n_________________"
    
    cell2 = signature_table.cell(0, 1)
    cell2.text = "SECOND PARTY:\n\n\n_________________"
    
    # Add witness section
    doc.add_paragraph("\n\n")
    witness_table = doc.add_table(rows=1, cols=2)
    cell1 = witness_table.cell(0, 0)
    cell1.text = "WITNESS 1:\n\n\n_________________"
    
    cell2 = witness_table.cell(0, 1)
    cell2.text = "WITNESS 2:\n\n\n_________________"
    
    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(
        temp_file.name,
        as_attachment=True,
        download_name=f"{contract_data.get('contract_type', 'Contract').replace(' ', '_')}_{contract_id}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def generate_tenancy_agreement(data):
    landlord = data.get('landlord', 'Landlord Name')
    tenant = data.get('tenant', 'Tenant Name')
    property_address = data.get('property_address', 'Property Address')
    rental_details = data.get('rental_details', 'Rental Details')
    security_deposit = data.get('security_deposit', 'Security Deposit Amount')
    
    # Generate a unique agreement number
    agreement_number = f"RLA-{uuid.uuid4().hex[:8].upper()}"
    
    # Get current date in Indian format
    today = datetime.datetime.now().strftime("%d-%m-%Y")
    
    draft = f"""
RENTAL LEASE AGREEMENT
(As per Indian Rent Control Act)

Agreement No: {agreement_number}
Date: {today}

THIS RENTAL AGREEMENT is made and executed at _______________ on this {today} between:

{landlord}, hereinafter referred to as the "LESSOR / LANDLORD" (which expression shall include his/her heirs, legal representatives, executors, and assigns) of the ONE PART

AND

{tenant}, hereinafter referred to as the "LESSEE / TENANT" (which expression shall include his/her heirs, legal representatives, executors, and assigns) of the OTHER PART.

WHEREAS the Lessor is the absolute owner and in possession of property situated at {property_address}, hereinafter referred to as the "SCHEDULED PREMISES".

AND WHEREAS the Lessee has approached the Lessor to let out the Scheduled Premises on rental basis, and the Lessor has agreed to let out the same on the following terms and conditions:

1. RENTAL AND LEASE PERIOD:
   The rent for the Scheduled Premises shall be {rental_details}. The minimum period of lease shall be 11 months starting from {today}.

2. SECURITY DEPOSIT:
   The Lessee has paid to the Lessor a sum of {security_deposit} as interest-free refundable security deposit, which shall be refunded at the time of vacating the premises after deducting any damages or dues if any.

3. PAYMENT OF RENT:
   The rent shall be paid in advance on or before the 10th day of each calendar month.

4. ELECTRICITY AND WATER CHARGES:
   The Lessee shall pay electricity and water charges based on actual consumption as per respective meter readings.

5. MAINTENANCE:
   General maintenance and minor repairs shall be the responsibility of the Lessee. Major structural repairs shall be the responsibility of the Lessor.

6. TERMINATION:
   Either party can terminate this agreement by giving one month's notice in writing.

7. RESTRICTIONS:
   The Lessee shall not sublet the premises or any part thereof to any person without prior written consent of the Lessor.

IN WITNESS WHEREOF, the parties hereto have set their hands on the day, month, and year first above written.

LESSOR / LANDLORD                                  LESSEE / TENANT
_________________                                  _________________
({landlord})                                       ({tenant})

WITNESS 1:                                         WITNESS 2:
_________________                                  _________________
    """.strip()
    return draft

def generate_contract(data):
    contract_type = data.get('contract_type', 'Contract Agreement')
    first_party = data.get('first_party', 'First Party Name and Address')
    second_party = data.get('second_party', 'Second Party Name and Address')
    contract_value = data.get('contract_value', 'Contract Value')
    duration = data.get('contract_duration', 'Contract Duration')
    terms = data.get('contract_terms', 'Specific Terms and Conditions')
    contract_id = data.get('contract_id', f"CNT-{uuid.uuid4().hex[:8].upper()}")
    
    # Get current date in Indian format
    today = datetime.datetime.now().strftime("%d-%m-%Y")
    
    draft = f"""
CONTRACT AGREEMENT - {contract_type.upper()}
(As per Indian Contract Act, 1872)

Contract No: {contract_id}
Date: {today}

THIS CONTRACT AGREEMENT is made and executed at _______________ on this {today} between:

{first_party}, hereinafter referred to as the "FIRST PARTY"

AND

{second_party}, hereinafter referred to as the "SECOND PARTY".

WHEREAS:
The parties have agreed to enter into a {contract_type} as detailed below.

NOW THIS AGREEMENT WITNESSES AS FOLLOWS:

1. SCOPE OF WORK:
   This {contract_type} covers the following services/goods/work as mutually agreed between the parties.

2. CONSIDERATION:
   The total consideration for this contract is {contract_value}.

3. TERM:
   This agreement shall be valid for a period of {duration}, commencing from {today}.

4. SPECIFIC TERMS AND CONDITIONS:
   {terms}

5. GOVERNING LAW:
   This agreement shall be governed by and construed in accordance with the laws of India, and the courts in India shall have exclusive jurisdiction.

6. DISPUTE RESOLUTION:
   Any dispute arising out of or in connection with this agreement shall first be resolved through amicable settlement. If not resolved within 30 days, the dispute shall be referred to arbitration in accordance with the Arbitration and Conciliation Act, 1996.

7. FORCE MAJEURE:
   Neither party shall be liable for any failure or delay in performance due to circumstances beyond its reasonable control.

IN WITNESS WHEREOF, the parties hereto have executed this agreement on the day and year first above written.

FIRST PARTY:                                     SECOND PARTY:
_________________                                _________________

WITNESS 1:                                       WITNESS 2:
_________________                                _________________
    """.strip()
    return draft

if __name__ == '__main__':
    app.run(debug=True)